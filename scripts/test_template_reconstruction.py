"""Smoke test — Phase 6 template reconstruction replaces stale content, keeps layout."""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document

from services.commentary_generator import generate_commentary
from services.template_formatter import format_enterprise_report

SAMPLE_RECORDS = [
    {
        "table": 1,
        "metric": "PAT",
        "period": "31.03.2025",
        "approved_value": 13720,
        "source_filename": "kotak mahindra BANK 25.pdf",
    },
    {
        "table": 1,
        "metric": "PAT",
        "period": "31.03.2024",
        "approved_value": 12000,
        "source_filename": "kotak mahindra BANK 25.pdf",
    },
    {
        "table": 1,
        "metric": "NII",
        "period": "31.03.2025",
        "approved_value": 25000,
        "source_filename": "kotak mahindra BANK 25.pdf",
    },
    {
        "table": 2,
        "metric": "PAT",
        "period": "H1FY26",
        "approved_value": 6535,
        "source_filename": "kotak mahindra BANK 25_3.pdf",
    },
    {
        "table": 2,
        "metric": "PAT",
        "period": "H1FY25",
        "approved_value": 6864,
        "source_filename": "kotak mahindra BANK 25_3.pdf",
    },
    {
        "table": 2,
        "metric": "NII",
        "period": "H1FY26",
        "approved_value": 14570,
        "source_filename": "kotak mahindra BANK 25_3.pdf",
    },
    {
        "table": 2,
        "metric": "NII",
        "period": "H1FY25",
        "approved_value": 13862,
        "source_filename": "kotak mahindra BANK 25_3.pdf",
    },
]


def _build_stale_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Credit Review Report")
    doc.add_paragraph("Old Template Bank Limited")
    doc.add_paragraph("As at 01 January 2020")
    doc.add_heading("Yearly Financials", level=1)
    yearly = doc.add_table(rows=3, cols=3)
    yearly.rows[0].cells[0].text = "Metric"
    yearly.rows[0].cells[1].text = "31.03.2025"
    yearly.rows[0].cells[2].text = "31.03.2024"
    yearly.rows[1].cells[0].text = "PAT"
    yearly.rows[1].cells[1].text = "2106"
    yearly.rows[1].cells[2].text = "1999"
    yearly.rows[2].cells[0].text = "NII"
    yearly.rows[2].cells[1].text = "9999"
    yearly.rows[2].cells[2].text = "8888"
    doc.add_heading("Half-Year Financials", level=1)
    half = doc.add_table(rows=3, cols=3)
    half.rows[0].cells[0].text = "Metric"
    half.rows[0].cells[1].text = "H1FY26"
    half.rows[0].cells[2].text = "H1FY25"
    half.rows[1].cells[0].text = "PAT"
    half.rows[1].cells[1].text = "1111"
    half.rows[1].cells[2].text = "2222"
    doc.add_heading("Profitability", level=2)
    doc.add_paragraph("PAT increased sharply to stale template values.")
    doc.add_heading("Capitalisation", level=2)
    doc.add_paragraph("Capitalisation is adequate with CAR at 99%.")
    doc.add_heading("Commentary", level=1)
    doc.add_paragraph("This is stale template commentary that must be removed.")
    doc.save(str(path))


def main() -> None:
    commentary = generate_commentary(SAMPLE_RECORDS)
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        template = tmp_path / "stale_template.docx"
        out_dir = tmp_path / "output"
        _build_stale_template(template)
        result = format_enterprise_report(
            template_path=template,
            reviewed_records=SAMPLE_RECORDS,
            commentary=commentary,
            warnings=[],
            output_dir=out_dir,
        )
        doc = Document(result["docx_path"])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )

        assert "2106" not in table_text, f"Stale PAT yearly value remained: {table_text}"
        assert "13720" in table_text, f"Expected extracted PAT not found: {table_text}"
        assert "stale template commentary" not in full_text.lower()
        assert "stale template values" not in full_text.lower()
        assert commentary["sections"][1]["paragraph"][:20] in full_text

        print("OK — template reconstruction replaced stale tables and narrative.")
        print("Log:", result.get("reconstruction_log", {}))


if __name__ == "__main__":
    main()
