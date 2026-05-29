"""
Optional analyst source map / methodology guidance file.

Guides interpretation methodology only — does not replace extracted financial values.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from services.pdf_reader import extract_pages_from_pdf

logger = logging.getLogger("credit_review")

SOURCE_MAP_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_SOURCE_MAP_CHARS = 200_000


@dataclass(frozen=True)
class SourceMapContext:
    """Parsed analyst guidance attached to the review session."""

    filename: str
    file_type: str
    text: str

    @property
    def char_count(self) -> int:
        return len(self.text)


def _truncate(text: str) -> str:
    if len(text) <= MAX_SOURCE_MAP_CHARS:
        return text
    return text[:MAX_SOURCE_MAP_CHARS] + "\n\n[… truncated for processing …]"


def _read_pdf_text(data: bytes, filename: str) -> str:
    pages = extract_pages_from_pdf(data, filename)
    return "\n\n".join(p.get("text", "") or "" for p in pages)


def _read_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _read_txt_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def load_source_map_from_bytes(filename: str, data: bytes) -> SourceMapContext | None:
    """Parse source map content from raw file bytes."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SOURCE_MAP_EXTENSIONS:
        logger.warning("Unsupported source map type: %s", suffix)
        return None
    if not data:
        return None

    try:
        if suffix == ".pdf":
            text = _read_pdf_text(data, filename)
        elif suffix == ".docx":
            text = _read_docx_text(data)
        else:
            text = _read_txt_text(data)
    except Exception as exc:
        logger.exception("Failed to read source map %s: %s", filename, exc)
        return None

    cleaned = _truncate((text or "").strip())
    if not cleaned:
        logger.warning("Source map %s contained no readable text", filename)
        return None

    return SourceMapContext(
        filename=filename,
        file_type=suffix.lstrip("."),
        text=cleaned,
    )


def load_source_map_from_upload(uploaded_file: Any) -> SourceMapContext | None:
    """Load from a Streamlit UploadedFile."""
    if uploaded_file is None:
        return None
    name = getattr(uploaded_file, "name", "") or "source_map"
    uploaded_file.seek(0)
    data = uploaded_file.getvalue()
    return load_source_map_from_bytes(name, data)


def methodology_summary(source_map: SourceMapContext | None, *, max_chars: int = 800) -> str:
    """Short excerpt for reports / commentary metadata."""
    if source_map is None or not source_map.text:
        return ""
    excerpt = source_map.text.replace("\n", " ").strip()
    if len(excerpt) > max_chars:
        excerpt = excerpt[:max_chars].rsplit(" ", 1)[0] + "…"
    return excerpt
